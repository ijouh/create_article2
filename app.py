import os
import time
import gspread
import openai
import language_tool_python
from flask import Flask, render_template, jsonify, Response, send_file, request
from google.oauth2.service_account import Credentials
from io import BytesIO
from markdown import markdown
from docx import Document
from wordpress_xmlrpc import Client, WordPressPost
from wordpress_xmlrpc.methods.posts import NewPost
from wordpress_xmlrpc.methods.media import UploadFile
from wordpress_xmlrpc.compat import xmlrpc_client
from werkzeug.utils import secure_filename

app = Flask(__name__)

# === CONFIGURATION MANUELLE ===
openai.api_key = os.getenv("GROQ_API_KEY")
openai.api_base = os.getenv("GROQ_API_BASE")

SHEET_ID = "169zrRGQ2tahhQrYIW147i82o_r66lOGuCdfeKpp0RXA"

WP_URL = os.getenv("WP_URL")
WP_USER = os.getenv("WP_USER")
WP_PASS = os.getenv("WP_PASS")


UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

tool = language_tool_python.LanguageToolPublicAPI('fr-FR')
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def get_google_sheet_data():
    creds = Credentials.from_service_account_file(
        "credentials.json",
        scopes=["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
    )
    client = gspread.authorize(creds)
    sheet = client.open_by_key(SHEET_ID).sheet1
    return sheet.get_all_records()


def build_prompt(sujet, lieu, mot_cle, mots_cles_secondaires, ton, public, longueur):
    return f"""
Tu es un rédacteur SEO expert. Rédige un article de blog de qualité professionnelle avec structure claire.

Sujet : {sujet} à {lieu}  
Mot-clé principal : {mot_cle}  
Mots-clés secondaires : {mots_cles_secondaires}  
Ton : {ton}  
Public cible : {public}  
Longueur souhaitée : {longueur} mots  

Structure l’article avec des titres (niveau 1 et 2), des paragraphes fluides, en style Markdown (avec `#` et `##`) mais sans balises HTML. Rends uniquement le texte de l’article.
    """


@app.route("/")
def home():
    return render_template("form.html")


@app.route("/list_rows")
def list_rows():
    data = get_google_sheet_data()
    return jsonify([{"index": i, "sujet": row.get("Sujet", "Sans sujet")} for i, row in enumerate(data)])


@app.route("/generate_from_sheet_stream/<int:row_index>")
def generate_from_sheet_stream(row_index):
    data = get_google_sheet_data()
    row = data[row_index]
    prompt = build_prompt(row["Sujet"], row["Lieu"], row["Mot_clé_principal"],
                          row["Mots_cles_secondaires"], row["Ton"],
                          row["Public_cible"], row["Longueur"])

    def stream():
        response = openai.ChatCompletion.create(
            model="llama3-70b-8192",
            messages=[{"role": "user", "content": prompt}],
            stream=True,
            temperature=0.4,
            max_tokens=4096,
        )
        for chunk in response:
            if "choices" in chunk:
                delta = chunk["choices"][0]["delta"]
                if "content" in delta:
                    time.sleep(0.01)
                    yield delta["content"]

    return Response(stream(), mimetype="text/plain")


@app.route("/download_article/<int:row_index>")
def download_article(row_index):
    data = get_google_sheet_data()
    row = data[row_index]
    titre = row.get("Sujet", "article")
    prompt = build_prompt(row["Sujet"], row["Lieu"], row["Mot_clé_principal"],
                          row["Mots_cles_secondaires"], row["Ton"],
                          row["Public_cible"], row["Longueur"])

    completion = openai.ChatCompletion.create(
        model="llama3-70b-8192",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
        max_tokens=4096,
    )

    texte = completion["choices"][0]["message"]["content"]
    texte_corrige = tool.correct(texte)

    doc = Document()
    for line in texte_corrige.splitlines():
        line = line.strip()
        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line:
            doc.add_paragraph(line)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"{titre.replace(' ', '_')}.docx"
    return send_file(file_stream, as_attachment=True,
                     download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/upload_image", methods=["POST"])
def upload_image():
    if 'image' not in request.files:
        return jsonify({"success": False, "error": "Pas de fichier image"}), 400
    file = request.files['image']
    if file.filename == '':
        return jsonify({"success": False, "error": "Fichier non sélectionné"}), 400
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filename = f"{int(time.time())}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        return jsonify({"success": True, "filename": filename})
    else:
        return jsonify({"success": False, "error": "Type de fichier non autorisé"}), 400


@app.route("/publish_to_wordpress/<int:row_index>", methods=["POST"])
def publish_to_wordpress(row_index):
    try:
        data = get_google_sheet_data()
        row = data[row_index]
        titre = row["Sujet"]
        prompt = build_prompt(row["Sujet"], row["Lieu"], row["Mot_clé_principal"],
                              row["Mots_cles_secondaires"], row["Ton"],
                              row["Public_cible"], row["Longueur"])

        completion = openai.ChatCompletion.create(
            model="llama3-70b-8192",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=4096,
        )

        texte = completion["choices"][0]["message"]["content"]
        texte_corrige = tool.correct(texte)

        wp = Client(WP_URL + 'xmlrpc.php', WP_USER, WP_PASS)

        image_url = f"https://source.unsplash.com/800x400/?{titre.replace(' ', '+')}"
        image_id = None

        if request.is_json and request.json.get("filename"):
            filename = request.json["filename"]
            image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.exists(image_path):
                with open(image_path, 'rb') as img:
                    data = {
                        'name': filename,
                        'type': 'image/jpeg',
                        'bits': xmlrpc_client.Binary(img.read())
                    }
                    response = wp.call(UploadFile(data))
                    image_url = response['url']
                    image_id = response['id']

        html_article = f'<img src="{image_url}" alt="{titre}" style="width:100%;height:auto;margin-bottom:20px;" />\n'
        html_article += markdown(texte_corrige)

        post = WordPressPost()
        post.title = titre
        post.content = html_article
        post.post_status = 'publish'

        # ✅ Utiliser la colonne "Catégories"
        categories = [cat.strip() for cat in row.get("Catégories", "Non classé").split(',')]
        post.terms_names = {'category': categories}

        if image_id:
            post.thumbnail = image_id

        post_id = wp.call(NewPost(post))
        return jsonify({"success": True, "message": f"✅ Article publié (ID: {post_id})"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(debug=True)
